"""Render a composed document plan (title + sections of typed blocks) to
Markdown, DOCX, and PDF. Pure rendering, no LLM calls -- placement was already
decided by compose.py; this just draws it.
"""

import shutil
from pathlib import Path
from xml.sax.saxutils import escape

from PIL import Image as PILImage


def _caption(block: dict, meta: dict) -> str:
    # compose.py's LLM sometimes leaves an image/table block's caption empty
    # despite the placement decision -- fall back to the classify-step caption,
    # which always exists, rather than rendering an empty caption.
    return block.get("caption") or meta.get("caption", "")


def render_markdown(title: str, sections: list[dict], images_by_id: dict, tables_by_id: dict, output_dir: Path) -> Path:
    images_dir = output_dir / "images"
    images_dir.mkdir(parents=True, exist_ok=True)

    lines = [f"# {title}", ""]
    for section in sections:
        lines.append(f"## {section['heading']}")
        lines.append("")
        for block in section["blocks"]:
            if block["type"] == "paragraph":
                if block["text"]:
                    lines.append(block["text"])
                    lines.append("")
            elif block["type"] == "image":
                image = images_by_id.get(block["ref"])
                if not image:
                    continue
                caption = _caption(block, image)
                image_name = f"{image['id']}.jpg"
                shutil.copy(image["path"], images_dir / image_name)
                lines.append(f"![{caption}](images/{image_name})")
                lines.append(f"*{caption}*")
                lines.append("")
            elif block["type"] == "table":
                table = tables_by_id.get(block["ref"])
                if not table:
                    continue
                lines.append("| " + " | ".join(table["headers"]) + " |")
                lines.append("| " + " | ".join("---" for _ in table["headers"]) + " |")
                for row in table["rows"]:
                    lines.append("| " + " | ".join(str(c) for c in row) + " |")
                lines.append(f"*{_caption(block, table)}*")
                lines.append("")

    output_path = output_dir / "document.md"
    output_path.write_text("\n".join(lines))
    return output_path


def render_docx(title: str, sections: list[dict], images_by_id: dict, tables_by_id: dict, output_path: Path) -> Path:
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading(title, level=0)

    for section in sections:
        doc.add_heading(section["heading"], level=1)
        for block in section["blocks"]:
            if block["type"] == "paragraph":
                if block["text"]:
                    doc.add_paragraph(block["text"])
            elif block["type"] == "image":
                image = images_by_id.get(block["ref"])
                if not image:
                    continue
                doc.add_picture(image["path"], width=Inches(5))
                doc.add_paragraph(_caption(block, image)).italic = True
            elif block["type"] == "table":
                table_data = tables_by_id.get(block["ref"])
                if not table_data:
                    continue
                rows = [table_data["headers"]] + table_data["rows"]
                table = doc.add_table(rows=len(rows), cols=len(table_data["headers"]))
                table.style = "Light Grid Accent 1"
                for r, row in enumerate(rows):
                    for c, val in enumerate(row):
                        table.cell(r, c).text = str(val)
                doc.add_paragraph(_caption(block, table_data)).italic = True

    doc.save(output_path)
    return output_path


def _image_dimensions(path: Path, max_width_inch: float = 4.5):
    from reportlab.lib.units import inch

    with PILImage.open(path) as img:
        w, h = img.size
    width = max_width_inch * inch
    height = width * (h / w)
    return width, height


def render_pdf(title: str, sections: list[dict], images_by_id: dict, tables_by_id: dict, output_path: Path) -> Path:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Image as RLImage
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    styles = getSampleStyleSheet()
    story = [Paragraph(escape(title), styles["Title"]), Spacer(1, 12)]

    for section in sections:
        story.append(Paragraph(escape(section["heading"]), styles["Heading1"]))
        for block in section["blocks"]:
            if block["type"] == "paragraph":
                if block["text"]:
                    story.append(Paragraph(escape(block["text"]), styles["BodyText"]))
                    story.append(Spacer(1, 6))
            elif block["type"] == "image":
                image = images_by_id.get(block["ref"])
                if not image:
                    continue
                width, height = _image_dimensions(Path(image["path"]))
                story.append(RLImage(image["path"], width=width, height=height))
                story.append(Paragraph(escape(_caption(block, image)), styles["Italic"]))
                story.append(Spacer(1, 10))
            elif block["type"] == "table":
                table_data = tables_by_id.get(block["ref"])
                if not table_data:
                    continue
                rows = [table_data["headers"]] + table_data["rows"]
                t = Table(rows)
                t.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#6d5ef8")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ]))
                story.append(t)
                story.append(Paragraph(escape(_caption(block, table_data)), styles["Italic"]))
                story.append(Spacer(1, 10))

    SimpleDocTemplate(str(output_path), pagesize=letter).build(story)
    return output_path
