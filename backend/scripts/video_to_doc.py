"""
video_to_doc.py
================
End-to-end pipeline: video file + transcript  ->  final PDF/DOCX document,
with charts/diagrams embedded as images and tables rebuilt as real structured tables.

    Video ──▶ [STAGE 0: extract + triage frames]  (local, free)
                        │
                        ▼
              [STAGE 1: classify + extract tables]   <── AI MODEL CALL #1 (vision)
                        │
    Transcript ─────────┤
                        ▼
              [STAGE 2: align into one timeline]  (local, free)
                        │
                        ▼
              [STAGE 3: compose document plan (JSON)]  <── AI MODEL CALL #2 (text/LLM)
                        │
                        ▼
              [STAGE 4: render to PDF / DOCX]  (local, free)

Only stages 1 and 3 need a paid/hosted model (or a self-hosted one). Everything else
runs locally for free. Search this file for "AI MODEL CALL" to find both integration points.

Requirements (system):    ffmpeg, tesseract-ocr
Requirements (python):    opencv-python-headless, pytesseract, numpy, reportlab, python-docx
See INSTRUCTIONS.md for full setup and how to wire in a real model.
"""

import os
import json
import subprocess
import numpy as np
import cv2
import pytesseract


# ===========================================================================
# STAGE 0 — Extract candidate frames from the video  (100% local, free)
# ===========================================================================
# No AI model involved anywhere in this stage. ffmpeg does extraction;
# tesseract + OpenCV do cheap heuristic triage to cut an hour of video down
# to a couple dozen frames actually worth showing an AI model.
# ===========================================================================

def _scene_change_frames(video_path: str, out_dir: str, scene_threshold: float) -> list[dict]:
    """Catches hard cuts / slide transitions."""
    pattern = os.path.join(out_dir, "scene_%04d.png")
    cmd = [
        "ffmpeg", "-y", "-i", video_path,
        "-vf", f"select='gt(scene,{scene_threshold})',showinfo",
        "-vsync", "vfr", pattern,
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)

    timestamps = []
    for line in result.stderr.splitlines():
        if "pts_time:" in line:
            timestamps.append(float(line.split("pts_time:")[1].split()[0]))

    frames = sorted(f for f in os.listdir(out_dir) if f.startswith("scene_"))
    return [{"path": os.path.join(out_dir, f), "timestamp": t}
            for f, t in zip(frames, timestamps)]


def _fixed_interval_frames(video_path: str, out_dir: str, interval_seconds: float) -> list[dict]:
    """
    Safety net: scene-change detection misses content that fades in gradually rather
    than cutting sharply (e.g. an animated chart). Dense fixed sampling catches those.
    """
    pattern = os.path.join(out_dir, "fixed_%04d.png")
    cmd = ["ffmpeg", "-y", "-i", video_path, "-vf", f"fps=1/{interval_seconds}", pattern]
    subprocess.run(cmd, capture_output=True, text=True)

    frames = sorted(f for f in os.listdir(out_dir) if f.startswith("fixed_"))
    return [{"path": os.path.join(out_dir, f), "timestamp": i * interval_seconds}
            for i, f in enumerate(frames)]


def extract_frames(video_path: str, out_dir: str,
                    scene_threshold: float = 0.15,
                    fixed_interval_seconds: float = 8.0) -> list[dict]:
    os.makedirs(out_dir, exist_ok=True)
    scene_frames = _scene_change_frames(video_path, out_dir, scene_threshold)
    fixed_frames = _fixed_interval_frames(video_path, out_dir, fixed_interval_seconds)
    return sorted(scene_frames + fixed_frames, key=lambda f: f["timestamp"])


def _phash(frame_path: str, size: int = 8) -> int:
    img = cv2.imread(frame_path, cv2.IMREAD_GRAYSCALE)
    small = cv2.resize(img, (size, size))
    avg = small.mean()
    bits = (small > avg).flatten()
    return sum(1 << i for i, b in enumerate(bits) if b)


def _hamming(a: int, b: int) -> int:
    return bin(a ^ b).count("1")


def _dedupe(hashes: list[int], threshold: int = 4) -> list[bool]:
    keep = [True] * len(hashes)
    for i in range(1, len(hashes)):
        if _hamming(hashes[i], hashes[i - 1]) <= threshold:
            keep[i] = False
    return keep


def _ocr_signal(frame_path: str) -> dict:
    text = pytesseract.image_to_string(frame_path).strip()
    words = [w for w in text.split() if len(w) > 1]
    return {"ocr_text": text, "word_count": len(words)}


def _chart_or_table_signal(frame_path: str) -> dict:
    img = cv2.imread(frame_path)
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    edges = cv2.Canny(gray, 50, 150)
    lines = cv2.HoughLinesP(edges, 1, np.pi / 180, threshold=120,
                             minLineLength=gray.shape[1] * 0.25, maxLineGap=10)
    num_long_lines = 0 if lines is None else len(lines)
    color_std = float(np.std(img.reshape(-1, 3), axis=0).mean())
    return {"num_long_lines": num_long_lines, "color_std": round(color_std, 1),
            "is_slide_like": num_long_lines >= 3 or color_std < 35}


def find_candidate_frames(video_path: str, out_dir: str,
                           scene_threshold: float = 0.15,
                           fixed_interval_seconds: float = 8.0) -> list[dict]:
    """Returns only the frames worth showing an AI model — the expensive step never
    has to look at 100% of the video."""
    frames = extract_frames(video_path, out_dir, scene_threshold, fixed_interval_seconds)

    hashes = [_phash(f["path"]) for f in frames]
    keep_flags = _dedupe(hashes)

    candidates = []
    for frame, keep in zip(frames, keep_flags):
        if not keep:
            continue
        ocr = _ocr_signal(frame["path"])
        cv_sig = _chart_or_table_signal(frame["path"])
        if ocr["word_count"] >= 3 or cv_sig["is_slide_like"]:
            candidates.append({**frame, **ocr, **cv_sig})

    return candidates


# ===========================================================================
# STAGE 1 — Classify each candidate frame + extract table structure
# ===========================================================================
#                         >>> AI MODEL CALL #1 <<<
#
# Needs a VISION-capable model (Claude, GPT-4o, Gemini Flash, or a self-hosted
# Qwen2.5-VL / DeepSeek-OCR). This is the one step Stage 0's cheap heuristics
# cannot do: deciding WHAT a frame shows and, for tables, turning the image
# into real rows/columns instead of a screenshot.
#
# `vision_model_call` is injected so you can swap providers without touching
# pipeline logic. It must take (image_path, prompt) and return the model's
# raw text response.
# ===========================================================================

CLASSIFY_PROMPT = """Look at this video frame and respond with ONLY valid JSON:

{{
  "type": "chart" | "table" | "diagram" | "photo" | "filler",
  "description": "one sentence describing what's shown, for use as a figure caption",
  "table": {{"headers": ["..."], "rows": [["..."]]}}   // only if type == "table", else null
}}

Context from cheap OCR on this frame (may be incomplete, use only as a hint): {ocr_hint}
"""


def classify_frame(frame: dict, vision_model_call) -> dict:
    """
    vision_model_call(image_path: str, prompt: str) -> str   (raw JSON text)

    See INSTRUCTIONS.md for example wiring against the Anthropic API.
    """
    prompt = CLASSIFY_PROMPT.format(ocr_hint=frame.get("ocr_text", "")[:300])
    raw = vision_model_call(frame["path"], prompt)
    result = json.loads(raw)
    result["timestamp"] = frame["timestamp"]
    result["path"] = frame["path"]
    return result


def build_visual_assets(candidates: list[dict], vision_model_call) -> tuple[list[dict], list[dict]]:
    """Splits classified frames into images_meta (charts/diagrams/photos worth embedding)
    and tables_meta (real structured data, keyed by timestamp)."""
    images_meta, tables_meta = [], []

    for frame in candidates:
        classified = classify_frame(frame, vision_model_call)
        if classified["type"] == "filler":
            continue
        elif classified["type"] == "table" and classified.get("table"):
            tables_meta.append({
                "timestamp": classified["timestamp"],
                "headers": classified["table"]["headers"],
                "rows": classified["table"]["rows"],
                "caption": classified["description"],
            })
        else:
            images_meta.append({
                "timestamp": classified["timestamp"],
                "path": classified["path"],
                "type": classified["type"],
                "caption": classified["description"],
            })

    return images_meta, tables_meta


# ===========================================================================
# STAGE 2 — Merge transcript + classified visuals into one timeline (local, free)
# ===========================================================================

def align_timeline(transcript_segments: list[dict], images_meta: list[dict],
                    tables_meta: list[dict]) -> list[dict]:
    """
    transcript_segments come from whatever transcription tool you use
    (Whisper, Otter, etc.) — that step is intentionally out of scope here.
    Expected shape: [{"start": 12.4, "text": "...", "speaker": "..."}, ...]
    """
    items = []
    for seg in transcript_segments:
        items.append({"timestamp": seg["start"], "kind": "transcript", "content": seg})
    for img in images_meta:
        items.append({"timestamp": img["timestamp"], "kind": "image", "content": img})
    for tbl in tables_meta:
        items.append({"timestamp": tbl["timestamp"], "kind": "table", "content": tbl})

    return sorted(items, key=lambda i: i["timestamp"])


# ===========================================================================
# STAGE 3 — Compose the document plan
# ===========================================================================
#                         >>> AI MODEL CALL #2 <<<
#
# Needs a TEXT-capable model (any modern LLM — this step is not vision, since
# stage 1 already turned visuals into descriptions/structured data). This is
# the step that reads the whole timeline and decides section structure,
# writes the prose, and places each image/table where it topically belongs
# (even if the transcript never explicitly mentions it at that timestamp).
#
# `text_model_call` takes a prompt string and returns raw text (JSON).
# ===========================================================================

COMPOSE_PROMPT = """You are given a chronological timeline of a video's transcript, \
extracted images, and extracted data tables. Write an organized document by TOPIC \
(not strictly chronological order) that explains the content in clear prose.

Return ONLY valid JSON matching this schema, nothing else:

{{
  "title": "string",
  "sections": [
    {{
      "heading": "string",
      "blocks": [
        {{"type": "paragraph", "text": "string"}},
        {{"type": "image", "ref": "<path from the timeline>", "caption": "string"}},
        {{"type": "table", "ref": "<table timestamp from the timeline>", "caption": "string"}}
      ]
    }}
  ]
}}

Rules:
- Place every image/table in the section it topically supports, even if the transcript
  never explicitly mentions it at that exact timestamp.
- Never invent table contents — only reference tables by timestamp; the renderer pulls
  the real structured data.

TIMELINE:
{timeline_json}
"""


def compose_document(timeline: list[dict], text_model_call) -> dict:
    """text_model_call(prompt: str) -> str  (raw JSON text)"""
    timeline_json = json.dumps(timeline, indent=2)
    prompt = COMPOSE_PROMPT.format(timeline_json=timeline_json)
    return json.loads(text_model_call(prompt))


# ===========================================================================
# STAGE 4 — Render the plan to PDF / DOCX (local, free)
# ===========================================================================

def render_to_pdf(plan: dict, tables_by_ts: dict, out_path: str):
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                     Image, Table, TableStyle)
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors

    styles = getSampleStyleSheet()
    story = [Paragraph(plan["title"], styles["Title"]), Spacer(1, 12)]

    for section in plan["sections"]:
        story.append(Paragraph(section["heading"], styles["Heading1"]))
        for block in section["blocks"]:
            if block["type"] == "paragraph":
                story.append(Paragraph(block["text"], styles["BodyText"]))
                story.append(Spacer(1, 6))
            elif block["type"] == "image":
                story.append(Image(block["ref"], width=4.5 * inch, height=2.5 * inch))
                story.append(Paragraph(block["caption"], styles["Italic"]))
                story.append(Spacer(1, 10))
            elif block["type"] == "table":
                tbl_data = tables_by_ts[block["ref"]]
                rows = [tbl_data["headers"]] + tbl_data["rows"]
                t = Table(rows)
                t.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1C2B4A")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ]))
                story.append(t)
                story.append(Paragraph(block["caption"], styles["Italic"]))
                story.append(Spacer(1, 10))

    SimpleDocTemplate(out_path, pagesize=letter).build(story)


def render_to_docx(plan: dict, tables_by_ts: dict, out_path: str):
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading(plan["title"], level=0)

    for section in plan["sections"]:
        doc.add_heading(section["heading"], level=1)
        for block in section["blocks"]:
            if block["type"] == "paragraph":
                doc.add_paragraph(block["text"])
            elif block["type"] == "image":
                doc.add_picture(block["ref"], width=Inches(4.5))
                doc.add_paragraph(block["caption"]).italic = True
            elif block["type"] == "table":
                tbl_data = tables_by_ts[block["ref"]]
                rows = [tbl_data["headers"]] + tbl_data["rows"]
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Light Grid Accent 1"
                for r, row in enumerate(rows):
                    for c, val in enumerate(row):
                        table.cell(r, c).text = str(val)
                doc.add_paragraph(block["caption"]).italic = True

    doc.save(out_path)


# ===========================================================================
# ORCHESTRATOR — runs all four stages end to end
# ===========================================================================

def run_pipeline(video_path: str,
                  transcript_segments: list[dict],
                  vision_model_call,          # AI MODEL CALL #1 — see Stage 1
                  text_model_call,            # AI MODEL CALL #2 — see Stage 3
                  work_dir: str = "/tmp/video_to_doc",
                  out_basename: str = "output",
                  formats: tuple = ("pdf", "docx")) -> dict:

    frames_dir = os.path.join(work_dir, "frames")
    candidates = find_candidate_frames(video_path, frames_dir)
    print(f"[stage 0] {len(candidates)} candidate frames selected for the vision model")

    images_meta, tables_meta = build_visual_assets(candidates, vision_model_call)
    print(f"[stage 1] classified into {len(images_meta)} images, {len(tables_meta)} tables")

    timeline = align_timeline(transcript_segments, images_meta, tables_meta)
    tables_by_ts = {t["timestamp"]: t for t in tables_meta}

    plan = compose_document(timeline, text_model_call)
    print(f"[stage 3] composed document plan with {len(plan['sections'])} sections")

    outputs = {}
    if "pdf" in formats:
        pdf_path = os.path.join(work_dir, f"{out_basename}.pdf")
        render_to_pdf(plan, tables_by_ts, pdf_path)
        outputs["pdf"] = pdf_path
    if "docx" in formats:
        docx_path = os.path.join(work_dir, f"{out_basename}.docx")
        render_to_docx(plan, tables_by_ts, docx_path)
        outputs["docx"] = docx_path

    print(f"[stage 4] rendered: {outputs}")
    return outputs


# ===========================================================================
# Example run with STUB model calls (for testing pipeline plumbing only —
# see INSTRUCTIONS.md to wire in a real model before using this for real)
# ===========================================================================

if __name__ == "__main__":
    import sys

    def stub_vision_model_call(image_path: str, prompt: str) -> str:
        # PLACEHOLDER ONLY. Replace with a real vision model call (see INSTRUCTIONS.md).
        # This stub just guesses from the OCR hint embedded in the prompt so the
        # pipeline can be smoke-tested without any API key.
        has_many_words = prompt.count(" ") > 40
        return json.dumps({
            "type": "table" if has_many_words else "diagram",
            "description": "Auto-generated placeholder caption (stub model).",
            "table": {"headers": ["Col A", "Col B"], "rows": [["1", "2"]]} if has_many_words else None,
        })

    def stub_text_model_call(prompt: str) -> str:
        # PLACEHOLDER ONLY. Replace with a real LLM call (see INSTRUCTIONS.md).
        return json.dumps({
            "title": "Stub Document (replace model calls for real output)",
            "sections": [{"heading": "Section 1",
                          "blocks": [{"type": "paragraph", "text": "Stub content."}]}],
        })

    video = sys.argv[1] if len(sys.argv) > 1 else "/mnt/user-data/uploads/apchem.mp4"
    fake_transcript = [{"start": 6.0, "text": "Intro", "speaker": "Speaker 1"}]

    run_pipeline(video, fake_transcript, stub_vision_model_call, stub_text_model_call,
                 work_dir="/home/claude/final_pipeline/work")
